"""
简历生成服务
支持基于优化建议生成新简历，并提供多种模板和导出格式
"""
import logging
from typing import Dict, Optional, List
from datetime import datetime
import json
import asyncio
from pathlib import Path
from playwright.async_api import async_playwright

from app.services.ai_service import ai_service
from app.db.session import SessionLocal
from app.models.resume import Resume


class ResumeGenerator:
    """简历生成器"""
    
    # 支持的模板 - 精美设计
    TEMPLATES = {
        "modern": {
            "name": "🌐 科技蓝",
            "description": "现代科技感十足，蓝色渐变主题，适合IT/互联网/技术岗位",
            "color_scheme": "blue"
        },
        "professional": {
            "name": "🏆 商务金",
            "description": "深邃商务蓝搭配金色点缀，高端大气，适合管理层/金融/商务岗位",
            "color_scheme": "gold"
        },
        "creative": {
            "name": "🎨 创意紫",
            "description": "梦幻紫色渐变，带有动态效果，适合设计师/创意/市场岗位",
            "color_scheme": "purple"
        },
        "minimal": {
            "name": "⚫ 极简黑",
            "description": "黑白简约风格，高级留白设计，适合高端岗位/各类场合",
            "color_scheme": "black"
        }
    }
    
    # 支持的导出格式
    EXPORT_FORMATS = ["pdf", "docx", "markdown", "html", "json", "png"]
    
    async def generate_optimized_resume(
        self,
        resume_id: str,
        job_id: Optional[str] = None,
        optimization_suggestions: Optional[Dict] = None,
        template: str = "modern",
        refined_content: Optional[str] = None
    ) -> Dict:
        """
        基于优化建议和目标职位生成深度优化简历
        """
        db = SessionLocal()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            if not resume or not resume.parsed_data:
                raise ValueError("简历不存在或未解析")
            
            job_data = None
            if job_id:
                from app.models.job import Job
                job = db.query(Job).filter(Job.id == job_id).first()
                if job:
                    job_data = job.parsed_data or {"title": job.title, "description": job.description}

            # 如果传入了建议，记录数量
            suggestions_count = 0
            if optimization_suggestions and "suggestions" in optimization_suggestions:
                suggestions_count = len(optimization_suggestions["suggestions"])
            elif not optimization_suggestions and job_data:
                # 如果没有显式传入建议但有目标职位，由 AI 判定为定向优化，设置一个虚拟计数或状态
                suggestions_count = 5 # 代表五大维度的标准深度优化

            # 调用更深度的 AI 内容生成
            optimized_content = await self._generate_resume_content_advanced(
                original_data=resume.parsed_data,
                job_data=job_data,
                suggestions=optimization_suggestions,
                template=template,
                refined_content=refined_content
            )
            
            # 生成纯文本版本供前端预览
            text_content = self._format_resume_as_text(optimized_content)
            
            return {
                "original_resume_id": resume_id,
                "target_job_id": job_id,
                "template": template,
                "template_info": self.TEMPLATES.get(template, self.TEMPLATES["modern"]),
                "content": optimized_content,
                "optimized_content": text_content,  # 新增：纯文本版本
                "metadata": {
                    "generated_at": datetime.utcnow().isoformat(),
                    "optimization_applied": True,
                    "suggestions_count": suggestions_count,
                    "target_job_title": job_data.get("title") if job_data else None,
                    "optimization_type": "针对性改写" if job_id else "通用提升"
                }
            }
        finally:
            db.close()

    async def _generate_resume_content_advanced(
        self,
        original_data: Dict,
        job_data: Optional[Dict],
        suggestions: Optional[Dict],
        template: str,
        refined_content: Optional[str] = None
    ) -> Dict:
        """
        高级 AI 简历内容生成引擎
        """
        target_context = f"目标职位：{json.dumps(job_data, ensure_ascii=False)}" if job_data else "通用职业发展优化"
        
        # 如果提供了用户修订版内容，则强制 AI 基于该内容进行结构化封装
        if refined_content:
            refinement_instruction = f"""
【核心指令：结构化用户修订稿】
用户已经对简历内容进行了手动修订，如下所示。你的任务是将其解析并填入简历的 JSON 结构中。
**绝对要求**：
1. 必须完全忠实于用户在【修订版文本】中提供的描述。
2. 将其拆解为 personal_info, work_experience, project_experience 等模块。
3. 确保 project_experience 中的每一个项目描述、行动和成果都源自用户的修订稿。
4. 仅在原稿中完全缺失的关键字段（如联系方式、教育背景）时，才从【原始数据】中补全。

【修订版文本】:
{refined_content}
"""
        else:
            refinement_instruction = f"""
【核心指令：深度改写建议应用】
重点应用以下改写建议：{json.dumps(suggestions, ensure_ascii=False)} 
执行全方位的深度内容增强，保持简历的真实性与专业度的平衡。
"""

        prompt = f"""
{refinement_instruction}
你是一位拥有15年经验的顶级职业顾问和 UI 视觉专家。请基于以下原始数据，为用户生成一份【极具视觉吸引力】且【内容深度优化】的简历。

【原始简历数据】:
{json.dumps(original_data, ensure_ascii=False, indent=2)}

【优化上下文】:
{target_context}

【改写建议】:
{json.dumps(suggestions, ensure_ascii=False, indent=2) if suggestions else "无特定建议"}

【⚠️ 简历生成黄金铁律 (STEEL RULES)】:
1. **内容继承承诺 (CONTENT INHERITANCE)**：原简历中的“项目名称”、“项目描述”和“职责细节”是受保护的资产，**绝对禁止删除、绝对禁止合并、绝对禁止用概括性的套话替换具体的事实信息。**
2. **润色而非改写 (POLISH, NOT REWRITE)**：你的角色是“抛光师”。如果原本的项目介绍写得已经很清楚，请原封不动地保留。你的优化仅限于：在保留原句的基础上，修正病句、提升话术专业度、或将口语化的描述改写为书面形式。
3. **新增成就点 (ADDITIONAL VALUE)**：你可以基于 JD 需求，为每个项目“额外增加” 1-2 条量化成果或技术动作点，但原有的点必须作为基石存在。
4. **格式对齐**：输出的 description 必须包含用户原有的项目背景介绍，actions 必须包含用户原有的全部技术动作，results 必须包含用户原有的全部成果。

请返回以下结构的 JSON 对象，确保 project_experience 数组内容充实：
{{
    "personal_info": {{
        "name": "姓名",
        "avatar_url": "{original_data.get('personal_info', {}).get('avatar_url', '')}",
        "title": "符合目标的专业职能头衔",
        "summary": "深度的职业画像（2-3句核心竞争力总结）",
        "labels": ["关键词1", "关键词2"],
        "contact": {{ "email": "...", "phone": "...", "location": "..." }}
    }},
    "work_experience": [
        {{
            "company": "...",
            "position": "...",
            "duration": "...",
            "description": "职责概况",
            "achievements": ["高价值成就1", "高价值成就2"]
        }}
    ],
    "project_experience": [
        {{
            "name": "项目名称",
            "role": "我的角色",
            "duration": "起止时间",
            "description": "项目面临的挑战与技术难度",
            "actions": ["我采取的关键技术方案1", "关键方案2"],
            "results": "最终实现的量化业务价值/技术指标"
        }}
    ],
    "skills_sections": [
        {{ "category": "技术领域", "skills": ["实打实的技能"] }}
    ],
    "education": [],
    "others": {{
        "certifications": ["证书1", "证书2"],
        "awards": ["奖项1", "奖项2"]
    }}
}}
{{
    "MATCH_WARNING": "请务必检查 project_experience 是否全量继承了原简历，严禁丢弃任何技术细节！"
}}
"""
        # 使用更大的 AI 限制或更专业的模型
        result = await ai_service._call_ai(prompt)
        return result or original_data

    def _generate_html_content(self, content: Dict, template_info: Dict) -> str:
        """
        生成极具视觉美感的现代 HTML 模板 (卡片化流式布局)
        """
        personal = content.get("personal_info", {})
        work_exp = content.get("work_experience", [])
        projects = content.get("project_experience", [])
        skills_sections = content.get("skills_sections", [])
        education = content.get("education", [])
        others = content.get("others", {})
        certifications = others.get("certifications", [])
        awards = others.get("awards", [])
        
        color = template_info.get("color_scheme", "blue")
        theme_colors = {
            "blue": "#1d4ed8",
            "gold": "#b45309",
            "purple": "#7e22ce",
            "black": "#0f172a"
        }
        primary_color = theme_colors.get(color, theme_colors["blue"])

        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{ size: a4; margin: 0; }}
        :root {{ 
            --primary: {primary_color};
            --bg-gray: #f8fafc;
            --text-dark: #1e293b;
            --text-muted: #64748b;
            --card-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }}
        body {{ 
            font-family: 'PingFang SC', 'HarmonyOS Sans', 'Microsoft YaHei', sans-serif; 
            background: var(--bg-gray); 
            padding: 0; margin: 0; color: var(--text-dark); 
            line-height: 1.5;
        }}
        
        /* 顶部 Banner */
        .resume-header {{ 
            background: linear-gradient(135deg, {primary_color} 0%, {primary_color}ee 100%);
            color: white; 
            padding: 50px 40px; 
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-bottom: 5px solid rgba(255,255,255,0.1);
        }}
        
        .header-main {{ flex: 1; }}
        .header-avatar {{
            width: 100px;
            height: 100px;
            border-radius: 20px;
            border: 4px solid rgba(255,255,255,0.2);
            background: white;
            overflow: hidden;
            margin-left: 40px;
            box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        }}
        .header-avatar img {{ width: 100%; height: 100%; object-fit: cover; }}
        
        .name-row {{ font-size: 34px; font-weight: 800; margin-bottom: 4px; letter-spacing: -0.5px; }}
        .title-row {{ font-size: 18px; opacity: 0.95; font-weight: 500; margin-bottom: 12px; }}
        .label-group {{ display: flex; gap: 8px; margin-bottom: 18px; }}
        .label-pill {{ background: rgba(255,255,255,0.15); padding: 3px 10px; border-radius: 6px; font-size: 11px; font-weight: 500; }}
        
        .contact-info {{ display: flex; gap: 20px; font-size: 13px; opacity: 0.85; }}
        
        .container {{ max-width: 960px; margin: -25px auto 40px; padding: 0 20px; }}
        
        /* 模块卡片化 */
        .section-card {{ 
            background: #fff; 
            border-radius: 16px; 
            padding: 24px; 
            margin-bottom: 20px; 
            box-shadow: var(--card-shadow);
        }}
        
        .section-title-box {{ 
            display: flex; 
            align-items: center; 
            gap: 10px; 
            margin-bottom: 20px; 
            padding-bottom: 12px;
            border-bottom: 2px solid #f1f5f9;
        }}
        .section-icon {{ font-size: 18px; }}
        .section-title {{ font-size: 17px; font-weight: 700; color: var(--text-dark); }}
        
        /* 经历项列表 */
        .item-row {{ margin-bottom: 24px; position: relative; padding-left: 18px; border-left: 2px solid #f1f5f9; }}
        .item-row::after {{ 
            content: ''; position: absolute; left: -5px; top: 6px; width: 8px; height: 8px; 
            background: #fff; border: 2px solid var(--primary); border-radius: 50%;
        }}
        
        .item-head {{ display: flex; justify-content: space-between; font-weight: 700; font-size: 14.5px; margin-bottom: 4px; }}
        .item-meta {{ font-size: 13.5px; color: var(--primary); font-weight: 600; margin-bottom: 10px; }}
        
        .bullet-list {{ margin: 0; padding: 0; list-style: none; }}
        .bullet-item {{ position: relative; padding-left: 15px; font-size: 13.5px; color: #444; margin-bottom: 6px; line-height: 1.6; }}
        .bullet-item::before {{ content: '•'; position: absolute; left: 0; color: var(--primary); font-weight: 800; }}
        
        /* 项目特别样式 */
        .project-block {{ background: #fbfcfe; padding: 16px; border-radius: 12px; border: 1px solid #edf2f7; margin-bottom: 15px; }}
        .project-result {{ margin-top: 10px; padding: 8px 12px; background: #f0fdf4; border-radius: 6px; font-size: 12.5px; color: #166534; font-weight: 500; }}
        
        /* 技能云 */
        .skill-group {{ margin-bottom: 12px; }}
        .skill-cat {{ font-size: 12px; font-weight: 700; color: var(--text-muted); text-transform: uppercase; margin-bottom: 8px; }}
        .skill-pills {{ display: flex; flex-wrap: wrap; gap: 6px; }}
        .skill-pill {{ background: #f1f5f9; color: #475569; padding: 4px 10px; border-radius: 6px; font-size: 11.5px; font-weight: 500; border: 1px solid #e2e8f0; }}

        .summary-box {{ background: #fff; padding: 24px; border-radius: 16px; margin-bottom: 20px; font-style: italic; font-size: 14.5px; line-height: 1.7; color: #4b5563; border-left: 4px solid var(--primary); }}
    </style>
</head>
<body>
    <div class="resume-card">
        <div class="resume-header">
            <div class="header-main">
                <div class="name-row">{personal.get('name', '姓名')}</div>
                <div class="title-row">{personal.get('title', '专业头衔')}</div>
                <div class="label-group">
                    {" ".join([f'<span class="label-pill">{L}</span>' for L in personal.get('labels', [])])}
                </div>
                <div class="contact-info">
                    <span>✉️ {personal.get('contact', {}).get('email', '-')}</span>
                    <span>📱 {personal.get('contact', {}).get('phone', '-')}</span>
                    <span>📍 {personal.get('contact', {}).get('location', '-')}</span>
                </div>
            </div>
            <div class="header-avatar">
                <img src="{personal.get('avatar_url') or 'https://ui-avatars.com/api/?name=' + personal.get('name', 'User') + '&background=random'}" alt="avatar">
            </div>
        </div>
        
        <div class="container">
            <div class="summary-box">“{personal.get('summary', '')}”</div>

            <div style="display: flex; gap: 20px;">
                <div style="flex: 2.5;">
                    <div class="section-card">
                        <div class="section-title-box">
                            <span class="section-icon">💼</span>
                            <span class="section-title">核心工作详细履历</span>
                        </div>
                        {self._render_work_exp_html(work_exp)}
                    </div>

                    <div class="section-card">
                        <div class="section-title-box">
                            <span class="section-icon">🚀</span>
                            <span class="section-title">核心项目深度解析</span>
                        </div>
                        {self._render_projects_html(projects)}
                    </div>
                </div>
                
                <div style="flex: 1;">
                    <div class="section-card">
                        <div class="section-title-box">
                            <span class="section-icon">🛠️</span>
                            <span class="section-title">专业技能栈</span>
                        </div>
                        {self._render_skills_html(skills_sections)}
                    </div>

                    <div class="section-card">
                        <div class="section-title-box">
                            <span class="section-icon">🎓</span>
                            <span class="section-title">教育背景</span>
                        </div>
                        {self._render_edu_html(education)}
                    </div>

                    {(certifications or awards) and f'''
                    <div class="section-card">
                        <div class="section-title-box">
                            <span class="section-icon">🏅</span>
                            <span class="section-title">荣誉认证</span>
                        </div>
                        {self._render_others_html(certifications, awards)}
                    </div>
                    '''}
                </div>
            </div>
        </div>
    </div>
</body>
</html>
"""
        return html

    def _render_work_exp_html(self, exps):
        html = ""
        for exp in exps:
            achievements = "".join([f'<li class="bullet-item">{a}</li>' for a in exp.get("achievements", [])])
            html += f"""
            <div class="item-row">
                <div class="item-head">
                    <span>{exp.get('company')}</span>
                    <span style="color: var(--text-muted); font-weight: 500;">{exp.get('duration')}</span>
                </div>
                <div class="item-meta">{exp.get('position')}</div>
                {f'<div style="font-size: 13px; color: #475569; margin-bottom: 8px;">{exp.get("description")}</div>' if exp.get("description") else ""}
                <ul class="bullet-list">{achievements}</ul>
            </div>
            """
        return html

    def _render_projects_html(self, projects):
        html = ""
        for p in projects:
            if not p: continue
            actions = "".join([f'<li class="bullet-item">{a}</li>' for a in p.get("actions", [])])
            html += f"""
            <div class="project-block">
                <div class="item-head">
                    <span>{p.get('name')}</span>
                    <span style="color: var(--text-muted); font-weight: 500; font-size: 12px;">{p.get('duration')}</span>
                </div>
                <div class="item-meta" style="margin-bottom: 8px;">{p.get('role')}</div>
                <div style="font-size: 13px; color: #475569; margin-bottom: 10px; font-weight: 500;">{p.get('description', '')}</div>
                <ul class="bullet-list">
                    {actions}
                </ul>
                <div class="project-result">
                    <strong>成果显著：</strong>{p.get('results', '')}
                </div>
            </div>
            """
        return html

    def _render_skills_html(self, sections):
        html = ""
        for s in sections:
            pills = "".join([f'<span class="skill-pill">{tag}</span>' for tag in s.get("skills", [])])
            html += f"""
            <div class="skill-group">
                <div class="skill-cat">{s.get('category')}</div>
                <div class="skill-pills">{pills}</div>
            </div>
            """
        return html

    def _render_edu_html(self, edu_list):
        html = ""
        for edu in edu_list:
            html += f"""
            <div style="margin-bottom: 12px;">
                <div style="font-weight: bold; font-size: 13.5px; color: var(--text-dark);">{edu.get('school')}</div>
                <div style="font-size: 12.5px; color: var(--primary); font-weight: 500;">{edu.get('degree')} · {edu.get('major')}</div>
                <div style="font-size: 11px; color: var(--text-muted);">{edu.get('duration')}</div>
            </div>
            """
        return html

    def _render_others_html(self, certifications, awards):
        html = ""
        if certifications:
            html += '<div class="skill-cat">资质证书</div><div class="skill-pills" style="margin-bottom: 15px;">'
            html += "".join([f'<span class="skill-pill" style="background:#fff7ed; color:#c2410c; border-color:#ffedd5;">{c}</span>' for c in certifications])
            html += '</div>'
        if awards:
            html += '<div class="skill-cat">核心荣誉</div><div class="skill-pills">'
            html += "".join([f'<span class="skill-pill" style="background:#fefce8; color:#a16207; border-color:#fef9c3;">{a}</span>' for a in awards])
            html += '</div>'
        return html
    async def export_resume(
        self,
        resume_data: Dict,
        format: str = "pdf",
        output_path: Optional[str] = None
    ) -> str:
        """
        导出简历到指定格式 (异步支持)
        """
        if format not in self.EXPORT_FORMATS:
            raise ValueError(f"不支持的格式: {format}")
        
        # 生成默认输出路径
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"exports/resume_{timestamp}.{format}"
        
        # 确保输出目录存在
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # 根据格式调用相应的导出方法
        if format == "json":
            return self._export_json(resume_data, output_path)
        elif format == "markdown":
            return self._export_markdown(resume_data, output_path)
        elif format == "html":
            return self._export_html(resume_data, output_path)
        elif format == "docx":
            return self._export_docx(resume_data, output_path)
        elif format == "pdf":
            return await self._export_pdf_advanced(resume_data, output_path)
        elif format == "png":
            return await self._export_image_advanced(resume_data, output_path)
        
        raise ValueError(f"格式 {format} 的导出功能尚未实现")

    def _export_json(self, resume_data: Dict, output_path: str) -> str:
        """导出为JSON格式"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(resume_data, f, ensure_ascii=False, indent=2)
        return output_path

    def _export_markdown(self, resume_data: Dict, output_path: str) -> str:
        """导出为Markdown格式"""
        content = resume_data.get("content", {})
        personal = content.get("personal_info", {})
        
        md = f"# {personal.get('name', '姓名')}\n\n"
        md += f"## {personal.get('title', '职位标题')}\n\n"
        md += f"{personal.get('summary', '')}\n\n"
        
        work_exp = content.get("work_experience", [])
        if work_exp:
            md += "## 工作经历\n\n"
            for exp in work_exp:
                md += f"### {exp.get('position')} @ {exp.get('company')}\n"
                md += f"*{exp.get('duration')}*\n\n"
                md += f"{exp.get('description', '')}\n\n"
                for ach in exp.get("achievements", []):
                    md += f"- {ach}\n"
                md += "\n"

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md)
        return output_path

    def _export_html(self, resume_data: Dict, output_path: str) -> str:
        """导出为HTML格式"""
        content = resume_data.get("content", {})
        template = resume_data.get("template", "modern")
        template_info = self.TEMPLATES.get(template, self.TEMPLATES["modern"])
        
        html_content = self._generate_html_content(content, template_info)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return output_path
    def _export_docx(self, resume_data: Dict, output_path: str) -> str:
        """导出为Word格式"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            content = resume_data.get("content", {})
            personal = content.get("personal_info", {})
            
            doc = Document()
            
            # 设置页边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # 标题
            title = doc.add_heading(personal.get('name', '姓名'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 副标题
            subtitle = doc.add_paragraph(personal.get('title', '职位标题'))
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # 个人简介
            if personal.get('summary'):
                summary = doc.add_paragraph(personal['summary'])
                summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
                summary_run = summary.runs[0]
                summary_run.font.size = Pt(11)
                summary_run.font.italic = True
            
            # 联系方式
            contact = personal.get("contact", {})
            contact_info = []
            if contact.get("email"):
                contact_info.append(f"📧 {contact['email']}")
            if contact.get("phone"):
                contact_info.append(f"📱 {contact['phone']}")
            if contact.get("location"):
                contact_info.append(f"📍 {contact['location']}")
            
            if contact_info:
                contact_para = doc.add_paragraph(" | ".join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_run = contact_para.runs[0]
                contact_run.font.size = Pt(10)
            
            doc.add_paragraph()  # 空行
            
            # 工作经历
            work_exp = content.get("work_experience", [])
            if work_exp:
                doc.add_heading('💼 工作经历', 1)
                
                for exp in work_exp:
                    # 职位和公司
                    position_para = doc.add_paragraph()
                    position_run = position_para.add_run(f"{exp.get('position', '职位')} @ {exp.get('company', '公司')}")
                    position_run.bold = True
                    position_run.font.size = Pt(12)
                    
                    # 时间
                    duration_para = doc.add_paragraph(exp.get('duration', '时间'))
                    duration_run = duration_para.runs[0]
                    duration_run.font.size = Pt(10)
                    duration_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    # 描述
                    if exp.get('description'):
                        doc.add_paragraph(exp['description'])
                    
                    # 亮点
                    highlights = exp.get("highlights", [])
                    if highlights:
                        for highlight in highlights:
                            doc.add_paragraph(highlight, style='List Bullet')
                    
                    doc.add_paragraph()  # 空行
            
            # 教育背景
            education = content.get("education", [])
            if education:
                doc.add_heading('🎓 教育背景', 1)
                
                for edu in education:
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"{edu.get('degree', '学位')} - {edu.get('major', '专业')}")
                    edu_run.bold = True
                    edu_run.font.size = Pt(12)
                    
                    school_para = doc.add_paragraph(f"{edu.get('school', '学校')} | {edu.get('duration', '时间')}")
                    school_run = school_para.runs[0]
                    school_run.font.size = Pt(10)
                    
                    doc.add_paragraph()  # 空行
            
            # 技能
            skills = content.get("skills", {})
            if skills:
                doc.add_heading('🛠️ 技能', 1)
                
                tech_skills = skills.get("technical", [])
                if tech_skills:
                    tech_para = doc.add_paragraph()
                    tech_para.add_run("技术技能：").bold = True
                    tech_para.add_run(", ".join(tech_skills))
                
                soft_skills = skills.get("soft", [])
                if soft_skills:
                    soft_para = doc.add_paragraph()
                    soft_para.add_run("软技能：").bold = True
                    soft_para.add_run(", ".join(soft_skills))
            
            # 保存文档
            doc.save(output_path)
            
            logging.info(f"简历已导出为Word: {output_path}")
            return output_path
            
        except ImportError:
            logging.error("python-docx 未安装，无法导出Word格式")
            raise ValueError("Word导出功能需要安装 python-docx 库")
    
    async def _export_pdf_advanced(self, resume_data: Dict, output_path: str) -> str:
        """
        使用 Playwright 生成高品质 PDF
        """
        content = resume_data.get("content", {})
        template = resume_data.get("template", "modern")
        template_info = self.TEMPLATES.get(template, self.TEMPLATES["modern"])
        
        html_content = self._generate_html_content(content, template_info)
        
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            
            # 设置 HTML 内容
            await page.set_content(html_content)
            # 等待网络空闲（如果有外部图片）
            await page.wait_for_load_state("networkidle")
            
            # 生成 PDF
            await page.pdf(
                path=output_path,
                format="A4",
                print_background=True,
                margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"}
            )
            
            await browser.close()
            
        return output_path

    async def _export_image_advanced(self, resume_data: Dict, output_path: str) -> str:
        """
        使用 Playwright 生成高清全景长图
        """
        content = resume_data.get("content", {})
        template = resume_data.get("template", "modern")
        template_info = self.TEMPLATES.get(template, self.TEMPLATES["modern"])
        
        html_content = self._generate_html_content(content, template_info)
        
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            # 设置较大的视口宽度，保持简历比例
            page = await browser.new_page(viewport={"width": 1000, "height": 1400}, device_scale_factor=2)
            
            await page.set_content(html_content)
            await page.wait_for_load_state("networkidle")
            
            # 获取页面真实高度
            height = await page.evaluate("document.body.scrollHeight")
            await page.set_viewport_size({"width": 1000, "height": height})
            
            # 截图
            await page.screenshot(path=output_path, full_page=True)
            
            await browser.close()
            
        return output_path
    
    def _format_resume_as_text(self, content: Dict) -> str:
        """
        将结构化简历数据格式化为易读的纯文本格式
        """
        lines = []
        
        # 个人信息
        personal = content.get("personal_info", {})
        if personal:
            lines.append("=" * 60)
            lines.append(f"  {personal.get('name', '未知姓名')}")
            if personal.get('title'):
                lines.append(f"  {personal.get('title')}")
            lines.append("=" * 60)
            lines.append("")
            
            # 联系方式
            contact = personal.get('contact', {})
            contact_info = []
            if contact.get('email'):
                contact_info.append(f"邮箱: {contact['email']}")
            if contact.get('phone'):
                contact_info.append(f"电话: {contact['phone']}")
            if contact.get('location'):
                contact_info.append(f"地址: {contact['location']}")
            if contact_info:
                lines.append(" | ".join(contact_info))
                lines.append("")
            
            # 职业概况
            if personal.get('summary'):
                lines.append("【职业概况】")
                lines.append(personal['summary'])
                lines.append("")
        
        # 工作经历
        work_exp = content.get("work_experience", [])
        if work_exp:
            lines.append("【工作经历】")
            lines.append("")
            for exp in work_exp:
                lines.append(f"▪ {exp.get('company', '未知公司')} | {exp.get('position', '未知职位')}")
                lines.append(f"  {exp.get('duration', '')}")
                if exp.get('description'):
                    lines.append(f"  {exp['description']}")
                
                achievements = exp.get('achievements', [])
                if achievements:
                    for achievement in achievements:
                        lines.append(f"  • {achievement}")
                lines.append("")
        
        # 项目经验
        projects = content.get("project_experience", [])
        if projects:
            lines.append("【项目经验】")
            lines.append("")
            for proj in projects:
                lines.append(f"▪ {proj.get('name', '未知项目')} | {proj.get('role', '未知角色')}")
                lines.append(f"  {proj.get('duration', '')}")
                if proj.get('description'):
                    lines.append(f"  {proj['description']}")
                
                actions = proj.get('actions', [])
                if actions:
                    for action in actions:
                        lines.append(f"  • {action}")
                
                if proj.get('results'):
                    lines.append(f"  成果: {proj['results']}")
                lines.append("")
        
        # 技能
        skills_sections = content.get("skills_sections", [])
        if skills_sections:
            lines.append("【专业技能】")
            lines.append("")
            for section in skills_sections:
                category = section.get('category', '技能')
                skills = section.get('skills', [])
                if skills:
                    lines.append(f"▪ {category}: {', '.join(skills)}")
            lines.append("")
        
        # 教育背景
        education = content.get("education", [])
        if education:
            lines.append("【教育背景】")
            lines.append("")
            for edu in education:
                school = edu.get('school', '未知学校')
                degree = edu.get('degree', '')
                major = edu.get('major', '')
                duration = edu.get('duration', '')
                lines.append(f"▪ {school} | {degree} {major}")
                if duration:
                    lines.append(f"  {duration}")
            lines.append("")
        
        # 荣誉认证
        others = content.get("others", {})
        cert_list = others.get("certifications", [])
        award_list = others.get("awards", [])
        
        # 兼容旧的 certifications 字段
        old_certs = content.get("certifications", [])
        if old_certs:
            cert_list.extend([c if isinstance(c, str) else c.get('name', '') for c in old_certs])
            # 去重
            cert_list = list(dict.fromkeys(cert_list))

        if cert_list or award_list:
            lines.append("【荣誉认证】")
            lines.append("")
            if cert_list:
                lines.append(f"▪ 资质证书: {', '.join(cert_list)}")
            if award_list:
                lines.append(f"▪ 荣誉奖项: {', '.join(award_list)}")
            lines.append("")
        
        return "\n".join(lines)


# 全局实例
resume_generator = ResumeGenerator()
