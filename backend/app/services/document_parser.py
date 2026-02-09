"""
文档解析服务
支持解析 PDF 和 Word 文档，提取职位描述文本
"""
import io
import logging
from typing import Optional
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParserService:
    """文档解析服务类"""
    
    @staticmethod
    async def parse_pdf(file_content: bytes) -> Optional[str]:
        """
        解析 PDF 文件，提取文本内容
        
        Args:
            file_content: PDF 文件的二进制内容
            
        Returns:
            提取的文本内容，失败返回 None
        """
        try:
            text_content = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            full_text = "\n".join(text_content).strip()
            
            if not full_text:
                logger.warning("PDF 文件中未提取到任何文本内容")
                return None
            
            logger.info(f"成功从 PDF 中提取 {len(full_text)} 个字符")
            return full_text
            
        except Exception as e:
            logger.error(f"解析 PDF 文件失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    @staticmethod
    async def parse_word(file_content: bytes) -> Optional[str]:
        """
        解析 Word 文档，提取文本内容
        
        Args:
            file_content: Word 文件的二进制内容
            
        Returns:
            提取的文本内容，失败返回 None
        """
        try:
            text_content = []
            
            doc = Document(io.BytesIO(file_content))
            
            # 提取所有段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = "\n".join(text_content).strip()
            
            if not full_text:
                logger.warning("Word 文档中未提取到任何文本内容")
                return None
            
            logger.info(f"成功从 Word 文档中提取 {len(full_text)} 个字符")
            return full_text
            
        except Exception as e:
            logger.error(f"解析 Word 文档失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    @staticmethod
    async def parse_document(file_content: bytes, filename: str) -> Optional[str]:
        """
        根据文件类型自动选择解析方法
        
        Args:
            file_content: 文件的二进制内容
            filename: 文件名（用于判断文件类型）
            
        Returns:
            提取的文本内容，失败返回 None
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return await DocumentParserService.parse_pdf(file_content)
        elif filename_lower.endswith(('.doc', '.docx')):
            return await DocumentParserService.parse_word(file_content)
        else:
            logger.error(f"不支持的文件类型: {filename}")
            return None


# 创建全局实例
document_parser = DocumentParserService()
